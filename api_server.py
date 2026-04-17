"""
碳策Agent - FastAPI 后端服务
包装现有 Python 服务为 REST API，不修改任何现有服务代码。
"""
import asyncio
import io
import os
import sys
import tempfile
import threading
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import requests

# 项目根目录加入 sys.path
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

from fastapi import FastAPI, HTTPException, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse

from agent.services.building_service import building_service
from agent.services.analysis_report_service import analysis_report_service
from agent.services.chat_history_service import chat_history_service
from agent.tools.building_enums import (
    BuildingType, Region, WallConstruction, RoofConstruction,
    RoofType, BuildingShape, HeatingType, CoolingType, WaterHeatingType,
    get_building_types, get_regions, get_wall_constructions,
    get_roof_constructions, get_roof_types, get_building_shapes,
    get_heating_types, get_cooling_types, get_water_heating_types,
)

# ---------- 线程锁 ----------
_building_lock = threading.Lock()

# ---------- FastAPI 实例 ----------
app = FastAPI(title="碳策Agent API", version="1.0.0")

# ---------- 挂载静态文件 ----------
web_dir = PROJECT_ROOT / "web"
app.mount("/static", StaticFiles(directory=str(web_dir)), name="static")


# ---------- Pydantic 模型 ----------

class BuildingCreate(BaseModel):
    name: str
    building_type: str = "办公建筑"
    region: str = "北方地区"
    floor_area_sqm: float = 1000.0
    num_floors: int = 1
    num_basements: int = 0
    year_built: Optional[int] = None
    num_employees: int = 0
    weekly_operating_hours: int = 40
    wall_construction: str = "混凝土"
    roof_construction: str = "混凝土平顶"
    roof_type: str = "平顶"
    building_shape: str = "矩形"
    glass_percentage: float = 30.0
    floor_to_ceiling_height_m: float = 3.0
    heating_type: str = "集中供暖"
    cooling_type: str = "中央空调"
    water_heating_type: str = "天然气热水器"
    uses_electricity: bool = True
    uses_natural_gas: bool = True


class BuildingUpdate(BaseModel):
    name: Optional[str] = None
    building_type: Optional[str] = None
    region: Optional[str] = None
    floor_area_sqm: Optional[float] = None
    num_floors: Optional[int] = None
    num_basements: Optional[int] = None
    year_built: Optional[int] = None
    num_employees: Optional[int] = None
    weekly_operating_hours: Optional[int] = None
    wall_construction: Optional[str] = None
    roof_construction: Optional[str] = None
    roof_type: Optional[str] = None
    building_shape: Optional[str] = None
    glass_percentage: Optional[float] = None
    floor_to_ceiling_height_m: Optional[float] = None
    heating_type: Optional[str] = None
    cooling_type: Optional[str] = None
    water_heating_type: Optional[str] = None
    uses_electricity: Optional[bool] = None
    uses_natural_gas: Optional[bool] = None


class EnergyDataAdd(BaseModel):
    month: str
    electricity_kwh: float = 0
    natural_gas_m3: float = 0
    water_m3: float = 0


class ChatMessage(BaseModel):
    role: str
    content: str


class ReportGenerateRequest(BaseModel):
    title: str = "建筑能耗分析报告"
    author: str = "碳策通智能体分析服务"
    building_ids: List[str] = []
    include_overview: bool = True
    include_area_analysis: bool = True
    include_type_analysis: bool = True
    include_energy_analysis: bool = True
    include_eui_analysis: bool = True
    include_recommendations: bool = True


# ---------- ReactAgent 单例 ----------
_agent_instance = None
_agent_lock = threading.Lock()


def get_agent():
    global _agent_instance
    if _agent_instance is None:
        with _agent_lock:
            if _agent_instance is None:
                from agent.react_agent import ReactAgent
                _agent_instance = ReactAgent()
    return _agent_instance


# ========== 枚举 API ==========

@app.get("/api/v1/enums/all")
async def get_all_enums():
    return {
        "building_types": get_building_types(),
        "regions": get_regions(),
        "wall_constructions": get_wall_constructions(),
        "roof_constructions": get_roof_constructions(),
        "roof_types": get_roof_types(),
        "building_shapes": get_building_shapes(),
        "heating_types": get_heating_types(),
        "cooling_types": get_cooling_types(),
        "water_heating_types": get_water_heating_types(),
    }


# ========== 建筑 API ==========

@app.get("/api/v1/buildings")
async def list_buildings(
    building_type: Optional[str] = None,
    region: Optional[str] = None,
    min_area: Optional[float] = None,
    max_area: Optional[float] = None,
    year_built_min: Optional[int] = None,
    year_built_max: Optional[int] = None,
):
    buildings = building_service.list_buildings(
        building_type=building_type,
        region=region,
        min_area=min_area,
        max_area=max_area,
        year_built_min=year_built_min,
        year_built_max=year_built_max,
    )
    return {"buildings": buildings, "total": len(buildings)}


@app.get("/api/v1/buildings/{building_id}")
async def get_building(building_id: str):
    building = building_service.get_building(building_id)
    if building is None:
        raise HTTPException(status_code=404, detail=f"建筑 {building_id} 不存在")
    return building


@app.post("/api/v1/buildings")
async def create_building(data: BuildingCreate):
    with _building_lock:
        result = building_service.create_building(
            name=data.name,
            building_type=data.building_type,
            region=data.region,
            floor_area_sqm=data.floor_area_sqm,
            num_floors=data.num_floors,
            num_basements=data.num_basements,
            year_built=data.year_built,
            num_employees=data.num_employees,
            weekly_operating_hours=data.weekly_operating_hours,
            wall_construction=data.wall_construction,
            roof_construction=data.roof_construction,
            roof_type=data.roof_type,
            building_shape=data.building_shape,
            glass_percentage=data.glass_percentage,
            floor_to_ceiling_height_m=data.floor_to_ceiling_height_m,
            heating_type=data.heating_type,
            cooling_type=data.cooling_type,
            water_heating_type=data.water_heating_type,
            uses_electricity=data.uses_electricity,
            uses_natural_gas=data.uses_natural_gas,
        )
    return result


@app.put("/api/v1/buildings/{building_id}")
async def update_building(building_id: str, data: BuildingUpdate):
    with _building_lock:
        kwargs = {k: v for k, v in data.model_dump().items() if v is not None}
        if not kwargs:
            raise HTTPException(status_code=400, detail="没有提供要更新的字段")
        result = building_service.update_building(building_id, **kwargs)
    if result is None:
        raise HTTPException(status_code=404, detail=f"建筑 {building_id} 不存在")
    return result


@app.delete("/api/v1/buildings/{building_id}")
async def delete_building(building_id: str):
    with _building_lock:
        success = building_service.delete_building(building_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"建筑 {building_id} 不存在")
    return {"message": f"建筑 {building_id} 已删除"}


@app.post("/api/v1/buildings/{building_id}/energy")
async def add_energy_data(building_id: str, data: EnergyDataAdd):
    with _building_lock:
        success = building_service.add_monthly_energy_data(
            building_id,
            month=data.month,
            electricity_kwh=data.electricity_kwh,
            natural_gas_m3=data.natural_gas_m3,
            water_m3=data.water_m3,
        )
    if not success:
        raise HTTPException(status_code=404, detail=f"建筑 {building_id} 不存在或数据添加失败")
    return {"message": "能耗数据添加成功"}


@app.get("/api/v1/buildings/{building_id}/energy")
async def get_energy_data(
    building_id: str,
    start_month: Optional[str] = None,
    end_month: Optional[str] = None,
):
    data = building_service.get_energy_data(building_id, start_month=start_month, end_month=end_month)
    return {"building_id": building_id, "energy_data": data}


# ========== 报告 API ==========

@app.get("/api/v1/reports")
async def list_reports():
    reports = analysis_report_service.list_reports()
    return {"reports": reports, "total": len(reports)}


@app.get("/api/v1/reports/{report_id}")
async def get_report(report_id: str):
    report = analysis_report_service.get_report(report_id)
    if report is None:
        raise HTTPException(status_code=404, detail=f"报告 {report_id} 不存在")
    return report


@app.get("/api/v1/reports/{report_id}/summary")
async def get_report_summary(report_id: str):
    summary = analysis_report_service.get_report_summary(report_id)
    if not summary:
        raise HTTPException(status_code=404, detail=f"报告 {report_id} 不存在")
    return {"report_id": report_id, "summary": summary}


@app.get("/api/v1/reports/search")
async def search_reports(q: str = Query(..., min_length=1)):
    results = analysis_report_service.search_reports(q)
    return {"query": q, "results": results, "total": len(results)}


# ========== 聊天 API ==========

@app.get("/api/v1/chat/history")
async def load_chat_history():
    messages = chat_history_service.load_history()
    info = chat_history_service.get_history_info()
    return {"messages": messages, "info": info}


@app.post("/api/v1/chat/history")
async def save_chat_history(data: List[ChatMessage]):
    messages = [{"role": m.role, "content": m.content} for m in data]
    success = chat_history_service.save_history(messages)
    return {"message": "聊天记录已保存" if success else "保存失败"}


@app.delete("/api/v1/chat/history")
async def clear_chat_history():
    success = chat_history_service.clear_history()
    return {"message": "聊天记录已清空" if success else "清空失败"}


@app.get("/api/v1/chat/stream")
async def chat_stream(q: str = Query(..., min_length=1)):
    """SSE 流式对话，桥接 ReactAgent.execute_stream()"""

    def event_generator():
        try:
            agent = get_agent()
            full_response = ""
            for chunk in agent.execute_stream(q):
                full_response += chunk
                yield {"event": "message", "data": chunk}
            # 发送完成事件，附带完整响应
            yield {"event": "done", "data": ""}
        except Exception as e:
            yield {"event": "error", "data": str(e)}

    return EventSourceResponse(event_generator())


# ========== 报告生成 API ==========

@app.post("/api/v1/report/generate")
async def generate_report(req: ReportGenerateRequest):
    """生成 Word 报告，返回文件下载"""
    try:
        import matplotlib
        if 'Agg' not in matplotlib.get_backend().lower():
            matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="需要安装 python-docx 和 matplotlib: pip install python-docx matplotlib")

    # 获取建筑数据
    all_buildings = building_service.list_buildings()
    if req.building_ids:
        buildings_data = {bid: all_buildings[bid] for bid in req.building_ids if bid in all_buildings}
    else:
        buildings_data = all_buildings

    if not buildings_data:
        raise HTTPException(status_code=400, detail="没有可分析的建筑数据")

    building_names = [b.get("basic_info", {}).get("name", "未命名") for b in buildings_data.values()]
    total_area = sum(b.get("basic_info", {}).get("floor_area_sqm", 0) for b in buildings_data.values())
    total_employees = sum(b.get("basic_info", {}).get("num_employees", 0) for b in buildings_data.values())

    temp_dir = tempfile.mkdtemp()
    try:
        doc = Document()

        # 标题
        title = doc.add_heading(req.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"报告作者: {req.author}")
        doc.add_paragraph(f"数据来源: 碳策通智能体分析服务")
        doc.add_paragraph()

        # 1. 建筑概览
        if req.include_overview:
            doc.add_heading("一、建筑概览", level=1)
            p = doc.add_paragraph()
            p.add_run(f"本次分析共涉及 {len(buildings_data)} 栋建筑，")
            p.add_run(f"总建筑面积 {total_area:,.0f} 平方米，")
            p.add_run(f"总员工数 {total_employees:,} 人。")
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "建筑ID", "建筑名称", "建筑面积(m²)", "建筑类型"
            for bid, b in buildings_data.items():
                row = table.add_row().cells
                row[0].text = bid
                row[1].text = b.get("basic_info", {}).get("name", "未命名")
                row[2].text = f"{b.get('basic_info', {}).get('floor_area_sqm', 0):,.0f}"
                row[3].text = b.get("basic_info", {}).get("building_type", "未知")

        # 2. 面积分析
        if req.include_area_analysis:
            doc.add_heading("二、建筑面积分析", level=1)
            areas = [b.get("basic_info", {}).get("floor_area_sqm", 0) for b in buildings_data.values()]
            avg_area = sum(areas) / len(areas) if areas else 0
            doc.add_paragraph(f"平均建筑面积: {avg_area:,.0f} 平方米")
            doc.add_paragraph(f"最大建筑面积: {max(areas) if areas else 0:,.0f} 平方米")
            doc.add_paragraph(f"最小建筑面积: {min(areas) if areas else 0:,.0f} 平方米")
            doc.add_paragraph()
            doc.add_paragraph("建筑面积分布图：")
            chart_path = os.path.join(temp_dir, "area_chart.png")
            _generate_area_chart(building_names, areas, chart_path)
            doc.add_picture(chart_path, width=Inches(5.5))

        # 3. 类型分析
        if req.include_type_analysis:
            doc.add_heading("三、建筑类型分布", level=1)
            type_counts = {}
            for b in buildings_data.values():
                bt = b.get("basic_info", {}).get("building_type", "未知")
                type_counts[bt] = type_counts.get(bt, 0) + 1
            for bt, count in type_counts.items():
                doc.add_paragraph(f"- {bt}: {count} 栋", style='List Bullet')
            if type_counts:
                doc.add_paragraph()
                doc.add_paragraph("建筑类型分布图：")
                chart_path = os.path.join(temp_dir, "type_chart.png")
                _generate_type_pie_chart(type_counts, chart_path)
                doc.add_picture(chart_path, width=Inches(4))

        # 4. 能耗分析
        if req.include_energy_analysis:
            doc.add_heading("四、能耗数据分析", level=1)
            energy_by_building = {}
            has_data = False
            for bid, b in buildings_data.items():
                name = b.get("basic_info", {}).get("name", bid)
                monthly = b.get("energy_consumption", {}).get("monthly_data", {})
                if monthly:
                    has_data = True
                    energy_by_building[name] = {
                        "electricity": sum(m.get("electricity_kwh", 0) for m in monthly.values()),
                        "gas": sum(m.get("natural_gas_m3", 0) for m in monthly.values()),
                        "water": sum(m.get("water_m3", 0) for m in monthly.values()),
                    }
            if has_data:
                te = sum(v["electricity"] for v in energy_by_building.values())
                tg = sum(v["gas"] for v in energy_by_building.values())
                tw = sum(v["water"] for v in energy_by_building.values())
                doc.add_paragraph(f"总用电量: {te:,.0f} kWh")
                doc.add_paragraph(f"总用气量: {tg:,.0f} m³")
                doc.add_paragraph(f"总用水量: {tw:,.0f} m³")
                doc.add_paragraph()
                doc.add_paragraph("建筑能耗对比图：")
                chart_path = os.path.join(temp_dir, "energy_chart.png")
                _generate_energy_chart(energy_by_building, chart_path)
                doc.add_picture(chart_path, width=Inches(5.5))
            else:
                doc.add_paragraph("暂无能耗数据")

        # 5. EUI 分析
        if req.include_eui_analysis:
            doc.add_heading("五、能耗强度(EUI)分析", level=1)
            eui_data = []
            for bid, b in buildings_data.items():
                name = b.get("basic_info", {}).get("name", bid)
                area = b.get("basic_info", {}).get("floor_area_sqm", 1)
                monthly = b.get("energy_consumption", {}).get("monthly_data", {})
                total_elec = sum(m.get("electricity_kwh", 0) for m in monthly.values())
                months = len(monthly) if monthly else 1
                eui = (total_elec / months * 12) / area if area > 0 else 0
                eui_data.append({"name": name, "eui": eui})
            if eui_data and any(d["eui"] > 0 for d in eui_data):
                for d in eui_data:
                    doc.add_paragraph(f"- {d['name']}: {d['eui']:.2f} kWh/m²/年", style='List Bullet')
                doc.add_paragraph()
                doc.add_paragraph("能耗强度对比图：")
                chart_path = os.path.join(temp_dir, "eui_chart.png")
                _generate_eui_chart(eui_data, chart_path)
                doc.add_picture(chart_path, width=Inches(5.5))
            else:
                doc.add_paragraph("暂无EUI数据")

        # 6. 智能建议
        if req.include_recommendations:
            doc.add_heading("六、智能节能建议", level=1)
            try:
                from model.factory import chat_model
                building_summary = []
                for bid, b in buildings_data.items():
                    basic = b.get("basic_info", {})
                    energy = b.get("energy_consumption", {}).get("monthly_data", {})
                    total_elec = sum(m.get("electricity_kwh", 0) for m in energy.values()) if energy else 0
                    area = basic.get("floor_area_sqm", 1)
                    eui = (total_elec / area * 12 / max(len(energy), 1)) if energy else 0
                    building_summary.append(
                        f"- {basic.get('name', bid)}: 类型={basic.get('building_type', '未知')}, "
                        f"面积={area:,.0f}m², 年份={basic.get('year_built', '未知')}, "
                        f"EUI={eui:.1f} kWh/m²/年, "
                        f"供暖={b.get('energy_systems', {}).get('heating', {}).get('primary_type', '未知')}, "
                        f"制冷={b.get('energy_systems', {}).get('cooling', {}).get('primary_type', '未知')}"
                    )
                prompt = (
                    f"请根据以下建筑数据，生成针对性的节能建议。每个建议包含：1.建议内容 2.建议理由 3.预计节能效果\n\n"
                    f"建筑数据:\n{''.join(building_summary)}\n\n"
                    f"请按【建议N】内容/理由/预计效果的格式输出5-8条建议。"
                )
                response = chat_model.invoke(prompt)
                ai_text = response.content if hasattr(response, 'content') else str(response)
                doc.add_paragraph("基于建筑特征和能耗数据分析，智能体为您提供以下个性化节能建议：")
                doc.add_paragraph()
                for line in ai_text.strip().split('\n'):
                    line = line.strip()
                    if line:
                        if line.startswith('内容：') or line.startswith('内容:'):
                            p = doc.add_paragraph()
                            r = p.add_run(line)
                            r.bold = True
                        else:
                            doc.add_paragraph(line)
            except Exception:
                # 备用建议
                doc.add_paragraph("基于建筑特征和能耗数据分析，为您提供以下节能建议：")
                doc.add_paragraph()
                fallbacks = [
                    "进行能耗审计，全面了解建筑能耗现状",
                    "更换LED照明系统，预计节能30-50%",
                    "安装建筑能源管理系统(EMCS)，综合节能10-15%",
                    "定期维护暖通空调设备，设备效率提升5-10%",
                    "采用变频空调系统，部分负荷效率提升20-30%",
                ]
                for fb in fallbacks:
                    doc.add_paragraph(f"- {fb}", style='List Bullet')

        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # 清理临时文件
        try:
            import shutil
            shutil.rmtree(temp_dir)
        except Exception:
            pass

        from urllib.parse import quote
        filename = f"{req.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except HTTPException:
        raise
    except Exception as e:
        # 清理临时文件
        try:
            import shutil
            if 'temp_dir' in dir():
                shutil.rmtree(temp_dir)
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"报告生成失败: {str(e)}")


# ========== 图表生成辅助函数 ==========

def _generate_area_chart(names, areas, output_path):
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12', '#9b59b6', '#1abc9c']
    bars = ax.bar(names, areas, color=[colors[i % len(colors)] for i in range(len(names))])
    ax.set_title('建筑面积分布', fontsize=14, fontweight='bold')
    ax.set_xlabel('建筑名称', fontsize=11)
    ax.set_ylabel('面积 (m2)', fontsize=11)
    ax.tick_params(axis='x', rotation=30)
    for bar, area in zip(bars, areas):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(areas) * 0.02,
                f'{area:,.0f}', ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def _generate_type_pie_chart(type_counts, output_path):
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(8, 8))
    colors = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12', '#9b59b6', '#1abc9c', '#e67e22', '#34495e']
    labels = list(type_counts.keys())
    values = list(type_counts.values())
    ax.pie(values, labels=labels, autopct='%1.1f%%',
           colors=[colors[i % len(colors)] for i in range(len(labels))],
           startangle=90, pctdistance=0.75)
    ax.set_title('建筑类型分布', fontsize=14, fontweight='bold')
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def _generate_energy_chart(energy_by_building, output_path):
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(12, 6))
    names = list(energy_by_building.keys())
    electricity = [v["electricity"] for v in energy_by_building.values()]
    gas = [v["gas"] for v in energy_by_building.values()]
    water = [v["water"] for v in energy_by_building.values()]
    x = range(len(names))
    width = 0.25
    ax.bar([i - width for i in x], electricity, width, label='电力(kWh)', color='#f1c40f')
    ax.bar([i for i in x], gas, width, label='天然气(m3)', color='#e74c3c')
    ax.bar([i + width for i in x], water, width, label='用水(m3)', color='#3498db')
    ax.set_title('建筑能耗对比', fontsize=14, fontweight='bold')
    ax.set_xlabel('建筑名称', fontsize=11)
    ax.set_ylabel('消耗量', fontsize=11)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=30, ha='right')
    ax.legend()
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def _generate_eui_chart(eui_data, output_path):
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 6))
    names = [d["name"] for d in eui_data]
    eui_values = [d["eui"] for d in eui_data]
    colors = ['#27ae60' if e < 100 else '#f39c12' if e < 200 else '#e74c3c' for e in eui_values]
    bars = ax.bar(names, eui_values, color=colors)
    ax.set_title('能耗强度对比 (EUI)', fontsize=14, fontweight='bold')
    ax.set_xlabel('建筑名称', fontsize=11)
    ax.set_ylabel('EUI (kWh/m2/年)', fontsize=11)
    ax.tick_params(axis='x', rotation=30)
    for bar, eui in zip(bars, eui_values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(eui_values) * 0.02,
                f'{eui:.1f}', ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


# ========== 碳排放监测 API (Carbon Monitor) ==========

CARBON_GLOBAL_URL = "https://datas.carbonmonitor.org/API/downloadFullDataset.php?source=carbon_global"
CARBON_CHINA_URL = "https://datas.carbonmonitor.org/API/downloadFullDataset.php?source=carbon_china"

# 内存缓存
_carbon_global_cache = {"data": None, "ts": 0}
_carbon_china_cache = {"data": None, "ts": 0}
CARBON_CACHE_TTL = 86400  # 缓存 24 小时

# 主要国家列表
CARBON_COUNTRIES = [
    "China", "United States of America", "India", "Germany",
    "Russian Federation", "Japan", "Brazil", "United Kingdom",
    "France", "Italy", "Canada", "Australia", "South Korea",
    "Indonesia", "Mexico", "Saudi Arabia", "South Africa",
    "Turkey", "Poland", "Thailand",
]

# 部门中英文映射
SECTOR_LABELS = {
    "Power": "电力",
    "Industry": "工业",
    "Ground Transport": "地面交通",
    "Residential": "居民生活",
    "Domestic Aviation": "国内航空",
    "International Aviation": "国际航空",
}

# 国家中文名
COUNTRY_LABELS = {
    "China": "中国",
    "United States of America": "美国",
    "India": "印度",
    "Germany": "德国",
    "Russian Federation": "俄罗斯",
    "Japan": "日本",
    "Brazil": "巴西",
    "United Kingdom": "英国",
    "France": "法国",
    "Italy": "意大利",
    "Canada": "加拿大",
    "Australia": "澳大利亚",
    "South Korea": "韩国",
    "Indonesia": "印尼",
    "Mexico": "墨西哥",
    "Saudi Arabia": "沙特阿拉伯",
    "South Africa": "南非",
    "Turkey": "土耳其",
    "Poland": "波兰",
    "Thailand": "泰国",
}

# 中国省份中文名映射
CHINA_PROVINCE_LABELS = {
    "Beijing": "北京", "Shanghai": "上海", "Tianjin": "天津",
    "Chongqing": "重庆", "Hebei": "河北", "Shanxi": "山西",
    "Inner Mongolia": "内蒙古", "Liaoning": "辽宁", "Jilin": "吉林",
    "Heilongjiang": "黑龙江", "Jiangsu": "江苏", "Zhejiang": "浙江",
    "Anhui": "安徽", "Fujian": "福建", "Jiangxi": "江西",
    "Shandong": "山东", "Henan": "河南", "Hubei": "湖北",
    "Hunan": "湖南", "Guangdong": "广东", "Guangxi": "广西",
    "Hainan": "海南", "Sichuan": "四川", "Guizhou": "贵州",
    "Yunnan": "云南", "Shaanxi": "陕西", "Gansu": "甘肃",
    "Qinghai": "青海", "Ningxia": "宁夏", "Xinjiang": "新疆",
    "Tibet": "西藏",
}


def _parse_carbon_csv(csv_text: str) -> list:
    """解析 Carbon Monitor CSV 为字典列表"""
    import csv
    from io import StringIO
    reader = csv.DictReader(StringIO(csv_text))
    rows = []
    for row in reader:
        try:
            val = float(row.get("value", 0))
        except (ValueError, TypeError):
            continue
        rows.append({
            "country": row.get("country", row.get("state", "")).strip(),
            "date": row.get("date", "").strip(),
            "sector": row.get("sector", "").strip(),
            "value": round(val, 4),
        })
    return rows


def _parse_date(date_str: str) -> str:
    """将 DD/MM/YYYY 格式转换为 YYYY-MM-DD"""
    parts = date_str.split("/")
    if len(parts) == 3:
        return f"{parts[2]}-{parts[1].zfill(2)}-{parts[0].zfill(2)}"
    return date_str


def _get_carbon_data(source: str) -> list:
    """获取并缓存碳排放 CSV 数据"""
    import time
    cache = _carbon_global_cache if source == "global" else _carbon_china_cache
    url = CARBON_GLOBAL_URL if source == "global" else CARBON_CHINA_URL

    now = time.time()
    if cache["data"] is not None and (now - cache["ts"]) < CARBON_CACHE_TTL:
        return cache["data"]

    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        rows = _parse_carbon_csv(resp.text)
        # 转换日期格式
        for r in rows:
            r["date"] = _parse_date(r["date"])
        cache["data"] = rows
        cache["ts"] = now
        print(f"  [Carbon Monitor] 已加载 {source} 数据: {len(rows)} 条记录")
        return rows
    except Exception as e:
        print(f"  [Carbon Monitor] 加载 {source} 数据失败: {e}")
        return cache["data"] or []


@app.get("/api/v1/carbon/global")
async def get_carbon_global(
    country: str = "",
    sector: str = "",
    year: str = "",
    start_date: str = "",
    end_date: str = "",
):
    """获取全球碳排放数据，支持按国家/部门/年份/日期范围过滤"""
    data = _get_carbon_data("global")
    if country:
        data = [r for r in data if r["country"] == country]
    if sector:
        data = [r for r in data if r["sector"] == sector]
    if year:
        data = [r for r in data if r["date"].startswith(year)]
    if start_date:
        data = [r for r in data if r["date"] >= start_date]
    if end_date:
        data = [r for r in data if r["date"] <= end_date]
    return {"success": True, "count": len(data), "data": data}


@app.get("/api/v1/carbon/china")
async def get_carbon_china(
    province: str = "",
    sector: str = "",
    year: str = "",
    start_date: str = "",
    end_date: str = "",
):
    """获取中国省级碳排放数据"""
    data = _get_carbon_data("china")
    if province:
        data = [r for r in data if r["country"] == province]
    if sector:
        data = [r for r in data if r["sector"] == sector]
    if year:
        data = [r for r in data if r["date"].startswith(year)]
    if start_date:
        data = [r for r in data if r["date"] >= start_date]
    if end_date:
        data = [r for r in data if r["date"] <= end_date]
    return {"success": True, "count": len(data), "data": data}


@app.get("/api/v1/carbon/summary")
async def get_carbon_summary(country: str = "", year: str = ""):
    """获取碳排放汇总数据（按国家和部门）"""
    data = _get_carbon_data("global")
    if country:
        data = [r for r in data if r["country"] == country]
    if year:
        data = [r for r in data if r["date"].startswith(year)]

    # 按国家汇总
    country_totals = {}
    for r in data:
        c = r["country"]
        if c not in country_totals:
            country_totals[c] = 0
        country_totals[c] += r["value"]

    # 按部门汇总
    sector_totals = {}
    for r in data:
        s = r["sector"]
        if s not in sector_totals:
            sector_totals[s] = 0
        sector_totals[s] += r["value"]

    # 获取日期范围
    dates = sorted(set(r["date"] for r in data))
    date_range = {"start": dates[0] if dates else "", "end": dates[-1] if dates else ""}

    return {
        "success": True,
        "date_range": date_range,
        "total_records": len(data),
        "countries": country_totals,
        "sectors": sector_totals,
    }


@app.get("/api/v1/carbon/trend")
async def get_carbon_trend(
    countries: str = "",
    sector: str = "",
    year: str = "",
    agg: str = "month",
):
    """
    获取碳排放趋势数据
    countries: 逗号分隔的国家列表，如 "China,United States of America"
    sector: 部门过滤
    year: 年份过滤
    agg: 聚合方式 - day/month/year
    """
    data = _get_carbon_data("global")

    country_list = [c.strip() for c in countries.split(",") if c.strip()] if countries else []
    if country_list:
        data = [r for r in data if r["country"] in country_list]
    if sector:
        data = [r for r in data if r["sector"] == sector]
    if year:
        data = [r for r in data if r["date"].startswith(year)]

    # 按时间聚合
    from collections import defaultdict
    trend = defaultdict(lambda: defaultdict(float))

    for r in data:
        if agg == "month":
            key = r["date"][:7]  # YYYY-MM
        elif agg == "year":
            key = r["date"][:4]  # YYYY
        else:
            key = r["date"]  # YYYY-MM-DD
        trend[key][r["country"]] += r["value"]

    # 转换为列表
    result = []
    for date_key in sorted(trend.keys()):
        item = {"date": date_key}
        item.update(trend[date_key])
        result.append(item)

    return {"success": True, "data": result, "agg": agg}


@app.get("/api/v1/carbon/countries")
async def get_carbon_countries():
    """获取可用的省份、部门、年份列表（中国数据）"""
    data = _get_carbon_data("china")
    provinces = sorted(set(r["country"] for r in data))
    sectors = sorted(set(r["sector"] for r in data))
    years = sorted(set(r["date"][:4] for r in data), reverse=True)
    return {
        "success": True,
        "provinces": provinces,
        "sectors": sectors,
        "years": years,
        "province_labels": CHINA_PROVINCE_LABELS,
        "sector_labels": SECTOR_LABELS,
    }


@app.get("/api/v1/carbon/china/summary")
async def get_carbon_china_summary(year: str = ""):
    """获取中国碳排放汇总数据（按省份和部门）"""
    data = _get_carbon_data("china")
    if year:
        data = [r for r in data if r["date"].startswith(year)]

    # 按省份汇总
    province_totals = {}
    for r in data:
        p = r["country"]
        if p not in province_totals:
            province_totals[p] = 0
        province_totals[p] += r["value"]

    # 按部门汇总
    sector_totals = {}
    for r in data:
        s = r["sector"]
        if s not in sector_totals:
            sector_totals[s] = 0
        sector_totals[s] += r["value"]

    dates = sorted(set(r["date"] for r in data))
    total = sum(province_totals.values())

    return {
        "success": True,
        "total": round(total, 2),
        "date_range": {"start": dates[0] if dates else "", "end": dates[-1] if dates else ""},
        "provinces": province_totals,
        "sectors": sector_totals,
        "province_count": len(province_totals),
    }


@app.get("/api/v1/carbon/china/trend")
async def get_carbon_china_trend(
    sector: str = "",
    year: str = "",
    agg: str = "month",
):
    """获取中国碳排放趋势数据（按时间聚合）"""
    data = _get_carbon_data("china")
    if sector:
        data = [r for r in data if r["sector"] == sector]
    if year:
        data = [r for r in data if r["date"].startswith(year)]

    from collections import defaultdict
    trend = defaultdict(lambda: {"total": 0.0})

    for r in data:
        if agg == "month":
            key = r["date"][:7]
        elif agg == "year":
            key = r["date"][:4]
        else:
            key = r["date"]
        trend[key]["total"] += r["value"]

    result = [{"date": k, "value": round(v["total"], 2)} for k, v in sorted(trend.items())]
    return {"success": True, "data": result, "agg": agg}


# ========== 根路由：返回 SPA ==========

@app.get("/")
async def serve_index():
    return FileResponse(str(web_dir / "index.html"))


# ========== 启动事件 ==========

@app.on_event("startup")
async def startup_event():
    print("=" * 50)
    print("  碳策Agent API 服务已启动")
    print("  访问 http://localhost:7899")
    print("=" * 50)
